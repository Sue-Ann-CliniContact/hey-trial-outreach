from docx import Document 
from datetime import date
import os
import io
import json
import openai
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# Set OpenAI API key from environment
openai.api_key = os.getenv("OPENAI_API_KEY")

# Your shared Google Drive folder ID
FOLDER_ID = "1ruHSgI3jo4rKKrLahitkNzFwGwT3yjCt"

def upload_to_drive(local_path, filename):
    creds_json = os.getenv("GOOGLE_CREDENTIALS")
    if not creds_json:
        raise ValueError("Missing GOOGLE_CREDENTIALS environment variable.")
    
    creds_dict = json.loads(creds_json)
    credentials = service_account.Credentials.from_service_account_info(
        creds_dict,
        scopes=["https://www.googleapis.com/auth/drive"]
    )

    service = build("drive", "v3", credentials=credentials)

    file_metadata = {
        "name": filename,
        "parents": [FOLDER_ID]
    }

    with open(local_path, "rb") as f:
        media = MediaIoBaseUpload(f, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        uploaded = service.files().create(body=file_metadata, media_body=media, fields="id").execute()

    file_id = uploaded.get("id")
    service.permissions().create(fileId=file_id, body={"role": "reader", "type": "anyone"}).execute()
    return f"https://drive.google.com/file/d/{file_id}/view?usp=sharing"

def generate_outreach_email(match, your_study_title, challenge_summary, success_summary="", agent_name="The CliniContact Team", agent_title="", output_folder="emails"):
    os.makedirs(output_folder, exist_ok=True)

    contact_name = match.get('contact_name', 'the research team')
    study_title = match.get('study_title') or match.get('title') or "your clinical study"

    prompt = f"""
You are a clinical outreach strategist at CliniContact.

Write a warm, intelligent, and personalized outreach email to {contact_name}, the study contact for:
"{study_title}"

You recently supported a study titled "{your_study_title}". The recruitment challenge was:
"{challenge_summary}".

{f"The results we achieved: {success_summary}" if success_summary else ""}

Your goal is to write a collegial and confident message introducing CliniContact. Mention that:
- We support research teams on complex recruitment and coordination, especially for underrepresented or hard-to-reach groups.
- We’re currently supporting similar studies and may have infrastructure, insights, or support relevant to this PI's work.
- You were impressed by their study’s focus (either condition, population, or methodology).
- Offer to connect for a brief call or conversation to explore how you might support their team.

Sign off as {agent_name}, {agent_title} from info@clinicontact.com.
Keep the tone warm, smart, and strategic—similar to how Kit from CliniContact writes outreach emails.
"""

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a thoughtful, persuasive outreach writer for a clinical recruitment company."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )

    email_text = response['choices'][0]['message']['content']

    doc = Document()
    doc.add_heading('Personalized Outreach Email', level=1)
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}\n")
    for line in email_text.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    filename = f"{match.get('nct_id', 'study')}_outreach.docx"
    local_path = os.path.join(output_folder, filename)
    doc.save(local_path)

    return upload_to_drive(local_path, filename)
